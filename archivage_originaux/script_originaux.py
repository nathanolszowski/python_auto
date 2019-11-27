#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""import"""

import sys
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

"""variables declarations"""

contract = {} #dict contains contracts
listContract = list() #list contains dicts of contracts

"""functions"""

def saisieNb(invite): #check if numbers
    while True:
        saisie= input(invite)
        try:
            saisie= int(saisie)
        except:
            print("Seul les caractères [0-9] sont autorisés.", file=sys.stderr)
        else:
            return saisie

"""core"""

#input contracts
invite="How many contracts do you want to insert ? : "
nbContract= saisieNb(invite)
i = 1
while i <= nbContract :
    print ('Insert contract n°', i)
    contract= {'year': input('Insert year : '),'month': input('Insert month : '), 'date': input('Insert date : '), 'counterpart' : input('Insert counterpart : '), 'title' : input('Insert title : '), 'object' : input('Insert object or external counterparts : '), 'negociator' : input('Insert negociator(s) : '), 'visa' : input('Insert visa (insert V if there is one else X) : '), 'actarus' : 'X', 'scanDj' : 'X', 'numFile' : 'X', 'original' : input('Insert original (insert V if there is one else X) : '), 'chest' : 'V'}
    listContract.append(contract)
    i += 1
print (listContract)

#data into excel file    
file= "inventaire.xlsx" #test if file already exist + do excel function
wb= load_workbook(file)
ws= wb['NJ']
colNum=1
while colNum<=3 :
    for key, value in contract.items():
        ws.cell(row=2, column=colNum).value = value
        colNum +=1
    wb.save(file)   

#enter data into new word file
document= Document() #test if file already exist

table = document.add_table(rows=0, cols=5)
table.style = 'Table Grid'
for contract in listContract:
    row_cells = table.add_row().cells
    row_cells[0].text = str(contract.get('year'))
    row_cells[1].text = str(contract.get('month')) 
    row_cells[2].text = str(contract.get('date'))
    row_cells[3].text = str(contract.get('counterpart'))
    row_cells[4].text = str(contract.get('title'))
    #paragraph space between tables
    user_address = document.add_paragraph()
    paragraph_format = document.styles['Normal'].paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_before.pt
    
document.save('etiquettes.docx')
